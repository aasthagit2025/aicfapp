from __future__ import annotations

import re
import tempfile
from collections import Counter
from pathlib import Path
from typing import BinaryIO, Dict, List, Tuple

import pandas as pd


def read_survey_file(uploaded_file: BinaryIO) -> Tuple[pd.DataFrame, Dict[str, str]]:
    name = uploaded_file.name.lower()
    labels: Dict[str, str] = {}

    if name.endswith(".csv"):
        return pd.read_csv(uploaded_file), labels

    if name.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded_file), labels

    if name.endswith(".sav"):
        try:
            import pyreadstat
        except ImportError as exc:
            raise RuntimeError("SPSS .sav support needs pyreadstat in requirements.txt.") from exc

        with tempfile.NamedTemporaryFile(delete=False, suffix=".sav") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        try:
            df, meta = pyreadstat.read_sav(tmp_path, apply_value_formats=False)
            labels = meta.column_names_to_labels or {}
            return df, labels
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    raise ValueError("Please upload a CSV, Excel, or SPSS .sav file.")


def extract_questionnaire_text(uploaded_file: BinaryIO | None) -> str:
    if uploaded_file is None:
        return ""

    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.getvalue().decode("utf-8", errors="ignore")

    if name.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX questionnaire support needs python-docx in requirements.txt.") from exc

        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name
        try:
            doc = Document(tmp_path)
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    return uploaded_file.getvalue().decode("utf-8", errors="ignore")


def readable_label(column: str, labels: Dict[str, str], questionnaire_text: str) -> str:
    if labels.get(column):
        return labels[column]

    cleaned = column.replace("_", " ").strip()
    if questionnaire_text and cleaned.lower() in questionnaire_text.lower():
        return cleaned
    return cleaned


def numeric_series(series: pd.Series) -> pd.Series:
    values = pd.to_numeric(series, errors="coerce").dropna()
    return values[values != 9]


def is_rating_1_to_5(values: pd.Series) -> bool:
    if values.empty:
        return False
    return values.min() >= 1 and values.max() <= 5 and values.nunique() >= 2


def is_nps_0_to_10(values: pd.Series) -> bool:
    if values.empty:
        return False
    return values.min() >= 0 and values.max() <= 10 and values.nunique() >= 5


def generate_rating_insight(insight_id: str, column: str, label: str, values: pd.Series) -> Dict[str, str]:
    mean = values.mean()
    top2 = (values >= 4).mean() * 100
    low = (values <= 2).mean() * 100

    if top2 >= 65 and low < 15:
        text = f"{label} appears to be a relative strength, with {top2:.1f}% top-two-box ratings and a mean score of {mean:.2f}/5."
    elif low >= 20 or top2 < 45:
        text = f"{label} needs human attention, as {low:.1f}% of valid respondents gave low ratings despite a mean score of {mean:.2f}/5."
    else:
        text = f"{label} shows moderate customer confidence, with a mean score of {mean:.2f}/5 and {top2:.1f}% top-two-box ratings."

    return {
        "insight_id": insight_id,
        "insight_text": text,
        "evidence_note": f"{column}: n={len(values)}, mean={mean:.2f}/5, top-two-box={top2:.1f}%, low ratings={low:.1f}%.",
    }


def generate_nps_insight(insight_id: str, column: str, label: str, values: pd.Series) -> Dict[str, str]:
    promoters = (values >= 9).mean() * 100
    passives = ((values >= 7) & (values <= 8)).mean() * 100
    detractors = (values <= 6).mean() * 100
    nps = promoters - detractors

    if nps >= 40 and detractors < 15:
        text = f"{label} shows strong customer advocacy, with an NPS-style score of {nps:.1f} and {promoters:.1f}% promoters."
    elif nps < 10 or detractors >= 30:
        text = f"{label} needs human attention because advocacy is weak, with an NPS-style score of {nps:.1f} and {detractors:.1f}% detractors."
    else:
        text = f"{label} shows moderate advocacy, with an NPS-style score of {nps:.1f}, {promoters:.1f}% promoters, {passives:.1f}% passives, and {detractors:.1f}% detractors."

    return {
        "insight_id": insight_id,
        "insight_text": text,
        "evidence_note": f"{column}: n={len(values)}, promoters={promoters:.1f}%, passives={passives:.1f}%, detractors={detractors:.1f}%, NPS-style score={nps:.1f}.",
    }


def generate_binary_insight(insight_id: str, column: str, label: str, values: pd.Series, total_n: int) -> Dict[str, str] | None:
    unique = set(values.dropna().unique().tolist())
    if not unique.issubset({0, 1}) and not unique.issubset({1, 2}):
        return None

    selected = (values == 1).sum()
    pct = selected / total_n * 100 if total_n else 0
    if pct < 10:
        return None

    return {
        "insight_id": insight_id,
        "insight_text": f"{label} is selected by {pct:.1f}% of respondents, making it a visible theme in the survey response pattern.",
        "evidence_note": f"{column}: selected n={int(selected)} out of total n={total_n}, selected percentage={pct:.1f}%.",
    }


def generate_text_insight(insight_id: str, column: str, label: str, series: pd.Series) -> Dict[str, str] | None:
    responses = [str(value).strip() for value in series.dropna() if str(value).strip()]
    if len(responses) < 10:
        return None

    stop_words = set(
        "the and for with this that have has had are was were you your our from product service services support customer customers solution solutions team very good great more need needs should can better improve improvement in on to of a an is it as by be we they their at all also".split()
    )
    words: List[str] = []
    for response in responses:
        for word in re.findall(r"[A-Za-z][A-Za-z-]{2,}", response.lower()):
            if word not in stop_words:
                words.append(word)

    top_words = [word for word, _ in Counter(words).most_common(6)]
    if not top_words:
        return None

    return {
        "insight_id": insight_id,
        "insight_text": f"Open-ended responses for {label} suggest recurring themes around {', '.join(top_words[:5])}; this should be coded qualitatively before client reporting.",
        "evidence_note": f"{column}: {len(responses)} open-ended responses reviewed; frequent terms include {', '.join(top_words)}.",
    }


def generate_insights(
    df: pd.DataFrame,
    labels: Dict[str, str] | None = None,
    questionnaire_text: str = "",
    max_insights: int = 25,
) -> pd.DataFrame:
    labels = labels or {}
    insights: List[Dict[str, str]] = []
    total_n = len(df)

    for column in df.columns:
        if len(insights) >= max_insights:
            break

        label = readable_label(str(column), labels, questionnaire_text)
        values = numeric_series(df[column])

        if len(values) >= max(10, total_n * 0.20):
            if is_rating_1_to_5(values):
                insights.append(generate_rating_insight(f"AI-{len(insights) + 1:03d}", str(column), label, values))
                continue
            if is_nps_0_to_10(values):
                insights.append(generate_nps_insight(f"AI-{len(insights) + 1:03d}", str(column), label, values))
                continue

            binary = generate_binary_insight(f"AI-{len(insights) + 1:03d}", str(column), label, values, total_n)
            if binary:
                insights.append(binary)
                continue

        if df[column].dtype == "object":
            text_insight = generate_text_insight(f"AI-{len(insights) + 1:03d}", str(column), label, df[column])
            if text_insight:
                insights.append(text_insight)

    return pd.DataFrame(insights)
